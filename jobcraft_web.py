import time
import os
import streamlit as st
import pandas as pd
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
import gspread 
import json 
import io 
from docx import Document 
from docx.shared import Pt 
from fpdf import FPDF 

# ---------------------------------------------------------
# 1. ESQUEMA DE DATOS (V4)
# ---------------------------------------------------------
class JobDescriptionV4(BaseModel):
    titulo_puesto: str = Field(description="El título que se mostrará en el perfil.")
    nivel: str = Field(description="Nivel de seniority del puesto.")
    titulo_oficial_match: str = Field(description="El título exacto encontrado en el catálogo oficial.")
    origen_titulo: str = Field(description="Debe decir 'ESTANDARIZADO' o 'NUEVO'.")
    mision_puesto: str = Field(description="Propósito principal del cargo.")
    responsabilidades_clave: list[str] = Field(description="5-7 funciones principales.")
    competencias_conductuales_seleccionadas: list[str] = Field(description="Las 4-5 competencias seleccionadas.")
    competencias_tecnicas: list[str] = Field(description="Habilidades duras.")
    requisitos_formacion: list[str] = Field(description="Formación académica.")
    kpis_sugeridos: list[str] = Field(description="KPIs.")
    observacion_ia: str = Field(description="Explicación de la equivalencia.")

GOOGLE_SHEET_ID = "1QPJ1JoCW7XO-6sf-WMz8SvAtylKTAShuMr_yGBoF-Xg" 

# ---------------------------------------------------------
# 2. CONEXIÓN A SHEETS
# ---------------------------------------------------------
def get_google_sheet_client():
    creds = st.secrets["gspread"]["gcp_service_account_credentials"]
    gc = gspread.service_account_from_dict(creds)
    return gc.open_by_key(GOOGLE_SHEET_ID)

@st.cache_data(ttl=3600)
def get_competencias(worksheet_name: str = "Diccionario_JobCraft"):
    try:
        sh = get_google_sheet_client()
        worksheet = sh.worksheet(worksheet_name)
        return pd.DataFrame(worksheet.get_all_records()), None
    except Exception as e:
        return None, f"Error cargando Diccionario: {e}"

@st.cache_data(ttl=3600)
def get_perfiles_estandar(worksheet_name: str = "Perfiles_Base_JobCraft"):
    try:
        sh = get_google_sheet_client()
        worksheet = sh.worksheet(worksheet_name)
        data = worksheet.get_all_records()
        lista_formateada = [f"{row['Cargo']} ({row.get('Nivel', 'N/A')})" for row in data]
        return "\n".join(lista_formateada), None
    except Exception as e:
        return "", f"Nota: No se encontró hoja de perfiles base ({e}). Se generará libremente."

# ---------------------------------------------------------
# 3. CEREBRO DE LA IA
# ---------------------------------------------------------
def run_jobcraft_ai(api_key: str, title: str, level: str, critical_skill: str, competencias_df: pd.DataFrame, lista_perfiles_base: str):
    max_retries = 3
    for attempt in range(max_retries):
        try:
            client = genai.Client(api_key=api_key)
            
            lista_competencias = "\n".join([
                f"- {row['Familia']}: {row['COREES_Definición_Core_N1_Inicial']}" 
                for index, row in competencias_df.iterrows()
            ])
            
            prompt = f"""
            Actúa como Director de Estructura Organizacional.
            Objetivo: Definir perfil para: '{title}' (Nivel: {level}).
            Habilidad Crítica: {critical_skill}
            
            --- CATÁLOGO OFICIAL DE PUESTOS ---
            {lista_perfiles_base}
            -----------------------------------
            
            INSTRUCCIONES DE ESTANDARIZACIÓN (HÍBRIDA):
            1. Busca en el CATÁLOGO OFICIAL si existe un puesto equivalente.
            2. SI ENCUENTRAS COINCIDENCIA:
               - 'titulo_puesto': Mantén el nombre que pidió el usuario.
               - 'titulo_oficial_match': Pon el nombre oficial del catálogo.
               - 'origen_titulo': "ESTANDARIZADO".
               - 'observacion_ia': "Este puesto es equivalente a [Titulo Oficial] en el Catálogo Maestro".
            3. SI NO HAY COINCIDENCIA:
               - 'titulo_puesto': El solicitado por el usuario.
               - 'titulo_oficial_match': "N/A"
               - 'origen_titulo': "NUEVO".
            
            INSTRUCCIONES DE CONTENIDO:
            4. Usa las competencias del diccionario adjunto.
            5. Genera Misión, Responsabilidades y KPIs profesionales.
            
            Genera JSON estricto.
            """
            
            config = types.GenerateContentConfig(response_mime_type="application/json", response_schema=JobDescriptionV4)
            response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt, config=config)
            return None, JobDescriptionV4(**json.loads(response.text))
            
        except Exception as e:
            if "503" in str(e) or "overloaded" in str(e).lower() or "429" in str(e):
                time.sleep(2)
                continue 
            else:
                return f"Error AI: {e}", None

    return "El servidor de IA está muy ocupado. Por favor intenta en unos segundos.", None

def generate_linkedin_post(api_key: str, job_data: JobDescriptionV4):
    try:
        client = genai.Client(api_key=api_key)
        prompt = f"""
        Escribe un POST DE LINKEDIN viral para:
        - Título: {job_data.titulo_puesto} ({job_data.nivel})
        - Misión: {job_data.mision_puesto}
        - Habilidades: {', '.join(job_data.competencias_tecnicas[:3])}
        
        Usa Emojis, estructura AIDA y hashtags.
        """
        response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
        return response.text
    except Exception as e:
        return f"No se pudo generar el post: {e}"

# ---------------------------------------------------------
# 4. GENERADORES DE ARCHIVOS (Word, PDF, TXT, CSV)
# ---------------------------------------------------------

# --- A. Generador de WORD ---
def create_docx(res):
    doc = Document()
    doc.add_heading(res.titulo_puesto, 0)
    p = doc.add_paragraph()
    p.add_run("Nivel: ").bold = True
    p.add_run(f"{getattr(res, 'nivel', 'N/A')}\n")
    p.add_run("Misión del Cargo: ").bold = True
    p.add_run(f"{res.mision_puesto}")

    def add_section(title, items):
        doc.add_heading(title, level=1)
        if items:
            for item in items:
                doc.add_paragraph(item, style='List Bullet')
    
    add_section("Responsabilidades Clave", res.responsabilidades_clave)
    add_section("Competencias Conductuales", res.competencias_conductuales_seleccionadas)
    add_section("Requisitos Técnicos", res.competencias_tecnicas)
    add_section("Formación", res.requisitos_formacion)
    add_section("KPIs Sugeridos", res.kpis_sugeridos)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- B. Generador de PDF ---
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'Perfil de Puesto - JobCraft AI', 0, 1, 'C')
        self.ln(5)

def clean_text_for_pdf(text):
    return text.encode('latin-1', 'replace').decode('latin-1')

def create_pdf(res):
    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    pdf.set_font("Arial", 'B', 16)
    pdf.multi_cell(0, 10, clean_text_for_pdf(res.titulo_puesto))
    pdf.ln(5)
    
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, "Misión del Cargo:", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 7, clean_text_for_pdf(res.mision_puesto))
    pdf.ln(5)
    
    def add_pdf_section(title, items):
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, clean_text_for_pdf(title), ln=True)
        pdf.set_font("Arial", size=12)
        for item in items:
            pdf.multi_cell(0, 7, f"- {clean_text_for_pdf(item)}")
        pdf.ln(3)

    add_pdf_section("Responsabilidades Clave", res.responsabilidades_clave)
    add_pdf_section("Competencias Conductuales", res.competencias_conductuales_seleccionadas)
    add_pdf_section("Requisitos Técnicos", res.competencias_tecnicas)
    add_pdf_section("Formación", res.requisitos_formacion)
    
    try:
        pdf_output = pdf.output(dest='S').encode('latin-1')
    except:
        pdf_output = pdf.output(dest='S')
    return pdf_output

# --- C. Generador de TXT ---
def convert_to_text(res):
    texto = f"PERFIL: {res.titulo_puesto}\nNivel: {getattr(res, 'nivel', 'N/A')}\nMisión: {res.mision_puesto}\n\n"
    texto += "RESPONSABILIDADES:\n" + "\n".join([f"- {x}" for x in res.responsabilidades_clave]) + "\n\n"
    texto += "COMPETENCIAS:\n" + "\n".join([f"- {x}" for x in res.competencias_conductuales_seleccionadas]) + "\n\n"
    texto += "TÉCNICAS:\n" + "\n".join([f"- {x}" for x in res.competencias_tecnicas])
    return texto

# --- D. Generador de CSV (NUEVO) ---
def create_csv(res):
    # Creamos un diccionario plano. Unimos las listas con " | "
    data = {
        "Título": res.titulo_puesto,
        "Nivel": getattr(res, 'nivel', 'N/A'),
        "Misión": res.mision_puesto,
        "Responsabilidades": " | ".join(res.responsabilidades_clave),
        "Competencias Conductuales": " | ".join(res.competencias_conductuales_seleccionadas),
        "Competencias Técnicas": " | ".join(res.competencias_tecnicas),
        "Formación": " | ".join(res.requisitos_formacion),
        "KPIs": " | ".join(res.kpis_sugeridos),
        "Estado": getattr(res, 'origen_titulo', 'N/A')
    }
    # Creamos un DataFrame de 1 sola fila
    df = pd.DataFrame([data])
    return df.to_csv(index=False).encode('utf-8')

# ---------------------------------------------------------
# 5. GUARDAR DATOS
# ---------------------------------------------------------
def guardar_datos_en_sheets(titulo_puesto: str, nivel: str, origen: str):
    try:
        sh = get_google_sheet_client()
        worksheet = sh.worksheet("Seguimiento Generaciones") 
        timestamp = pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')
        worksheet.append_row([timestamp, titulo_puesto, nivel, origen]) 
        return True, None
    except Exception as e:
        return False, f"Error al guardar: {e}"

# ---------------------------------------------------------
# 6. INTERFAZ GRÁFICA
# ---------------------------------------------------------
st.set_page_config(page_title="JobCraft AI Suite", layout="wide", page_icon="👔") 

st.markdown("## 👔 JobCraft AI: Suite de Reclutamiento")
st.markdown("---")

api_key = st.secrets["GEMINI_API_KEY"] if "GEMINI_API_KEY" in st.secrets else None
if not api_key:
    st.error("⚠️ Falta API KEY en Secrets")
    st.stop()

col_load1, col_load2 = st.columns(2)
with col_load1:
    df_comp, err_comp = get_competencias()
    if err_comp: st.error(err_comp); st.stop()
    st.success(f"✅ Diccionario: {len(df_comp)} registros", icon="📘")

with col_load2:
    txt_perfiles, err_perf = get_perfiles_estandar()
    if "Error" in str(err_perf): 
        st.warning(err_perf)
    else:
        st.success(f"✅ Catálogo Oficial conectado", icon="🗂️")

with st.container():
    col1, col2, col3 = st.columns([1, 1, 2])
    with col1:
        t = st.text_input("Nombre del Cargo (Búsqueda)", value="Analista de Ventas")
    with col2:
        l = st.selectbox("Nivel de Seniority", ["Junior (0-2 años)", "Semi-Senior (3-5 años)", "Senior (5+ años)", "Líder/Gerente"])
    with col3:
        s = st.text_input("Habilidad Crítica / Foco", placeholder="Ej: Python, Ventas B2B...")

    btn = st.button("✨ Generar Perfil Técnico", type="primary", use_container_width=True)

if btn:
    st.session_state['job_result'] = None 
    with st.spinner("🔍 Diseñando perfil..."):
        err_ai, res = run_jobcraft_ai(api_key, t, l, s, df_comp, txt_perfiles)
        
        if err_ai: 
            st.error(err_ai)
        else:
            st.session_state['job_result'] = res
            origen_seguro = getattr(res, 'origen_titulo', 'NUEVO')
            nivel_seguro = getattr(res, 'nivel', l)
            guardar_datos_en_sheets(res.titulo_puesto, nivel_seguro, origen_seguro)

# --- VISUALIZACIÓN DE RESULTADOS ---
if 'job_result' in st.session_state and st.session_state['job_result']:
    res = st.session_state['job_result']
    
    st.divider()
    
    # --- ÁREA DE DESCARGAS (Ahora con CSV) ---
    st.subheader("📂 Exportar Perfil")
    
    # Creamos 4 columnas para que quepan todos los botones
    col_d1, col_d2, col_d3, col_d4 = st.columns(4)
    
    file_name_base = f"Perfil_{res.titulo_puesto.replace(' ', '_')}"
    
    with col_d1:
        docx_file = create_docx(res)
        st.download_button(
            label="📄 Word (.docx)",
            data=docx_file,
            file_name=f"{file_name_base}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    with col_d2:
        try:
            pdf_bytes = create_pdf(res)
            st.download_button(
                label="📕 PDF (.pdf)",
                data=bytes(pdf_bytes),
                file_name=f"{file_name_base}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        except Exception as e:
            st.error("Error PDF")

    with col_d3:
        txt_file = convert_to_text(res)
        st.download_button(
            label="📝 Texto (.txt)",
            data=txt_file,
            file_name=f"{file_name_base}.txt",
            mime="text/plain",
            use_container_width=True
        )
        
    with col_d4:
        # Botón CSV Nuevo
        csv_file = create_csv(res)
        st.download_button(
            label="📊 Datos (.csv)",
            data=csv_file,
            file_name=f"{file_name_base}.csv",
            mime="text/csv",
            use_container_width=True
        )

    st.divider()

    # --- VISUALIZACIÓN EN PANTALLA ---
    if getattr(res, 'origen_titulo', 'NUEVO') == "ESTANDARIZADO":
        st.success(f"✅ PUESTO VALIDADO EN CATÁLOGO")
    else:
        st.info(f"🆕 NUEVO PUESTO CREADO")
        
    st.markdown(f"<h1 style='text-align: center; color: #1E88E5;'>{res.titulo_puesto}</h1>", unsafe_allow_html=True)

    titulo_oficial = getattr(res, 'titulo_oficial_match', 'N/A')
    obs_ia = getattr(res, 'observacion_ia', '')
    
    if titulo_oficial != "N/A" and titulo_oficial != res.titulo_puesto:
            st.warning(f"⚠️ **Nota:** Oficialmente equivale a **'{titulo_oficial}'**.")
    elif obs_ia:
            st.caption(f"🤖 Nota: {obs_ia}")

    st.markdown(f"<p style='text-align: center;'>Nivel: <b>{getattr(res, 'nivel', 'N/A')}</b></p>", unsafe_allow_html=True)
    st.info(f"🎯 **Misión:** {res.mision_puesto}")
    
    col_izq, col_der = st.columns(2)
    with col_izq:
        st.subheader("🚀 Responsabilidades")
        for item in res.responsabilidades_clave: st.markdown(f"✅ {item}")
        st.subheader("🧠 Competencias (ADN)")
        for item in res.competencias_conductuales_seleccionadas: st.markdown(f"🔹 {item}")
    with col_der:
        st.subheader("🛠️ Técnicas")
        for item in res.competencias_tecnicas: st.markdown(f"🔧 {item}")
        st.subheader("🎓 Requisitos")
        for item in res.requisitos_formacion: st.markdown(f"🎓 {item}")
    
    st.divider()
    
    st.markdown("### 📢 Modo Reclutador")
    if st.button("🚀 Generar Post para LinkedIn"):
        with st.spinner("✍️ Redactando anuncio viral..."):
            post_linkedin = generate_linkedin_post(api_key, res)
            st.markdown("#### 📝 Tu Post sugerido:")
            st.text_area("Copia este texto:", value=post_linkedin, height=300)
            st.balloons()
